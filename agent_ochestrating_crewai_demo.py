import os
from dotenv import load_dotenv
from datetime import datetime
from crewai import Agent, Task, Crew, Process
from langchain_google_genai import ChatGoogleGenerativeAI
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
from google.api_core.exceptions import ResourceExhausted, TooManyRequests
from docx import Document
from docx.shared import Inches

load_dotenv()

# ============== LLM with Rate Limit Protection ==============
@retry(
    stop=stop_after_attempt(5),
    wait=wait_exponential(multiplier=1, min=8, max=60),
    retry=retry_if_exception_type((ResourceExhausted, TooManyRequests)),
    reraise=True
)
def create_llm():
    return ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",      # Best for free tier (higher limits)
        temperature=0.65,
        google_api_key=os.getenv("GEMINI_API_KEY"),
        convert_system_message_to_human=True,
    )

llm = create_llm()
print("✅ Using Google Gemini 2.5 Flash with retry protection")

# ==================== AGENTS ====================
process_designer = Agent(
    role="Principal Engineering Manager",
    goal="Design a professional 4-round interview process for SDE-3 MERN Tech Lead",
    backstory="You have 10+ years hiring senior engineers at FAANG-level companies.",
    verbose=True,
    llm=llm
)

mern_expert = Agent(
    role="Staff MERN Stack Engineer",
    goal="Create deep, practical SDE-3 level technical questions on MERN stack",
    backstory="You have built and scaled large production MERN applications serving millions of users.",
    verbose=True,
    llm=llm
)

system_design_expert = Agent(
    role="Principal System Design Architect",
    goal="Create relevant, real-world system design questions suitable for Tech Lead",
    backstory="Expert in designing highly scalable distributed systems.",
    verbose=True,
    llm=llm
)

evaluator = Agent(
    role="Senior Tech Lead & Evaluator",
    goal="Compile everything into a professional interview kit with model answers and scoring rubric. Also suggest simple diagram descriptions where helpful.",
    backstory="You have conducted 500+ interviews and create clear, shareable evaluation kits.",
    verbose=True,
    llm=llm
)

# ==================== TASKS ====================
task1 = Task(
    description="Design a complete 4-round interview process for SDE-3 MERN Stack Tech Lead role. Include duration, goal, and who should interview in each round.",
    agent=process_designer,
    expected_output="Clean structured 4-round interview process"
)

task2 = Task(
    description="Create 12 strong SDE-3 level technical questions covering React, Node.js, MongoDB, Express, performance, security, and scaling.",
    agent=mern_expert,
    expected_output="Categorized numbered list of questions"
)

task3 = Task(
    description="Create 2 excellent system design questions for Tech Lead level (e.g. scalable feed, checkout system). Include key discussion points and trade-offs.",
    agent=system_design_expert,
    expected_output="2 detailed system design questions with guidance"
)

task4 = Task(
    description="""Create ONE beautiful, professional final interview kit in markdown format.
    Title: SDE-3 MERN Stack Tech Lead Interview Kit
    Include:
    - Interview Process table
    - All technical questions (categorized)
    - 2 System Design questions with discussion points
    - Model answers for at least 6 important questions
    - Scoring rubric (1-10 scale)
    - For important topics (e.g. architecture, scaling, feed system), suggest 1-2 simple diagram descriptions (e.g., "High-level diagram of fan-out write model") that can be turned into images.""",
    agent=evaluator,
    expected_output="Complete professional markdown interview kit",
    context=[task1, task2, task3]
)

# ==================== CREW ====================
timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

crew = Crew(
    agents=[process_designer, mern_expert, system_design_expert, evaluator],
    tasks=[task1, task2, task3, task4],
    process=Process.sequential,
    verbose=True,           # Logs go only to the log file, not polluting terminal much
    memory=True,
    output_log_file=f"sde3_mern_interview_logs_{timestamp}.log"
)

result = crew.kickoff()

# Save Markdown
md_filename = f"SDE3_MERN_Interview_Kit_{timestamp}.md"
with open(md_filename, "w", encoding="utf-8") as f:
    f.write(str(result))

print(f"✅ Markdown saved: {md_filename}")

# ==================== Create Professional DOCX ====================
doc = Document()
doc.add_heading('SDE-3 MERN Stack Tech Lead Interview Kit', 0)
doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

# Add the full content
doc.add_paragraph(str(result))

# Optional: Add placeholder for images (you can replace later)
doc.add_heading('Suggested Diagrams', level=2)
doc.add_paragraph("1. High-level architecture diagram for scalable feed system (Fan-out vs Fan-in)")
doc.add_paragraph("2. MongoDB schema + indexing strategy for social feed")
doc.add_paragraph("3. Horizontal scaling diagram for Node.js backend")

docx_filename = f"SDE3_MERN_Interview_Kit_{timestamp}.docx"
doc.save(docx_filename)
print(f"✅ Professional DOCX saved: {docx_filename}")

print("\n🎉 All done! Files are ready for GitHub and LinkedIn.")
print("   - Use the .md file for GitHub README")
print("   - Use the .docx file for sharing / printing")
print("   - Generate images from the suggested diagram descriptions and insert them in the DOCX.")